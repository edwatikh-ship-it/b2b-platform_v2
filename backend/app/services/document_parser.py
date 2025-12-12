import os
import json
import logging
from typing import Dict, List, Optional
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
from openpyxl import load_workbook

logger = logging.getLogger(__name__)


class DocumentParser:
    def __init__(self):
        """Initialize document parser with Groq client"""
        self.groq_api_key = os.getenv("GROQ_API_KEY")
        self.client = None
        
        if self.groq_api_key:
            try:
                from groq import Groq
                self.client = Groq(api_key=self.groq_api_key)
                logger.info("[Parser] Groq API Key: ✅ SET")
            except Exception as e:
                logger.error(f"[Parser] Failed to init Groq: {e}")
                logger.info("[Parser] Groq API Key: ❌ NOT SET")
        else:
            logger.info("[Parser] Groq API Key: ❌ NOT SET")

    async def parse_document(self, file_path: str) -> Dict:
        """Parse document and extract positions"""
        logger.info("[Parser] ===== PARSING START =====")
        logger.info(f"[Parser] File: {file_path}")
        
        try:
            # Extract raw text
            raw_text = self._extract_text(file_path)
            logger.info(f"[Extract] Reading file: {file_path}")
            logger.info(f"[Parser] Extracted text length: {len(raw_text)} chars")
            logger.info(f"[Parser] First 200 chars:\n{raw_text[:200]}")
            
            if not raw_text or len(raw_text.strip()) < 10:
                logger.warning("[Parser] Text too short for parsing")
                return {
                    "positions": [],
                    "metadata": {"confidence": 0, "method": "empty"},
                    "preview": raw_text,
                    "raw_text": raw_text
                }
            
            # Try Groq first
            if self.client:
                logger.info("[Parser] 🚀 Attempting Groq/Llama parsing...")
                result = await self._parse_with_groq(raw_text)
                if result["positions"]:
                    logger.info(f"[Parser] ✅ Groq SUCCESS: {len(result['positions'])} positions")
                    return result
                else:
                    logger.warning("[Parser] ⚠️  Groq returned empty, falling back to regex")
            
            # Fallback to regex
            logger.info("[Parser] 📋 Using REGEX fallback...")
            result = self._parse_with_regex(raw_text)
            logger.info(f"[Parser] Regex found {len(result['positions'])} positions")
            result["preview"] = raw_text
            result["raw_text"] = raw_text
            return result
            
        except Exception as e:
            logger.error(f"[Parser] FATAL ERROR: {e}")
            import traceback
            traceback.print_exc()
            return {
                "positions": [],
                "metadata": {"confidence": 0, "method": "error", "error": str(e)},
                "preview": "",
                "raw_text": ""
            }

    def _extract_text(self, file_path: str) -> str:
        """Extract text from DOCX, PDF, or XLSX (including tables!)"""
        try:
            ext = file_path.lower().split(".")[-1]
            
            if ext == "docx":
                doc = DocxDocument(file_path)
                
                # Сначала параграфы
                text = "\n".join([p.text for p in doc.paragraphs])
                
                # ПОТОМ ТАБЛИЦЫ! (это было пропущено!)
                for table in doc.tables:
                    text += "\n"
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        text += "\n" + row_text
                
                logger.info(f"[Extract] DOCX: Extracted {len(text)} chars (paragraphs + tables)")
                return text
            
            elif ext == "pdf":
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                logger.info(f"[Extract] PDF: Extracted {len(text)} chars from {len(reader.pages)} pages")
                return text
            
            elif ext == "xlsx":
                wb = load_workbook(file_path)
                text = ""
                for ws in wb.sheetnames:
                    sheet = wb[ws]
                    for row in sheet.iter_rows(values_only=True):
                        text += " ".join([str(cell) if cell else "" for cell in row]) + "\n"
                logger.info(f"[Extract] XLSX: Extracted {len(text)} chars")
                return text
            
            else:
                logger.warning(f"[Extract] Unknown format: .{ext}")
                return ""
                
        except Exception as e:
            logger.error(f"[Extract] ERROR: {e}")
            return ""

    async def _parse_with_groq(self, text: str) -> Dict:
        """Parse using Groq Llama 3.3 (latest model)"""
        try:
            logger.info("[Groq] Starting Llama 3.3 parsing...")
            logger.info("[Groq] Sending request to Groq API...")
            
            prompt = """Проанализируй этот документ и извлеки список позиций для закупки.

Формат ответа - JSON массив объектов:
{
  "positions": [
    {"pos": 1, "name": "название товара", "unit": "ед.измерения", "qty": количество},
    ...
  ]
}

ВАЖНО:
- Заполни ТОЛЬКО валидный JSON
- Укажи реальные номера позиций, названия, единицы измерения (м, шт, кг, л и т.д.) и количество
- Если позиции на русском - оставь как есть
- Для unit используй: м, шт, кг, л, м2, м3, комплект, упаковка и т.д.

Текст документа:
""" + text

            response = self.client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.3,
                max_tokens=1024,
            )
            
            logger.info("[Groq] Response received")
            response_text = response.choices[0].message.content
            logger.info(f"[Groq] Response: {len(response_text)} chars")
            
            # Parse JSON from response
            json_str = response_text
            if "```json" in response_text:
                json_str = response_text.split("```json")[1].split("```")[0]
            elif "```" in response_text:
                json_str = response_text.split("```")[1].split("```")[0]
            
            data = json.loads(json_str.strip())
            positions = data.get("positions", [])
            
            logger.info(f"[Groq] Parsed {len(positions)} positions")
            
            return {
                "positions": positions,
                "metadata": {
                    "confidence": 85,
                    "method": "groq_llama",
                    "model": "llama-3.3-70b-versatile"
                },
                "preview": self._format_preview(positions),
                "raw_text": text
            }
            
        except json.JSONDecodeError as e:
            logger.error(f"[Groq] JSON parse error: {e}")
            return {"positions": [], "metadata": {"confidence": 0, "method": "groq_json_error"}}
        
        except AttributeError as e:
            logger.error(f"[Groq] API attribute error: {e}")
            logger.error("[Groq] Full traceback:")
            import traceback
            logger.error(traceback.format_exc())
            return {"positions": [], "metadata": {"confidence": 0, "method": "groq_api_error"}}
        
        except Exception as e:
            logger.error(f"[Groq] ERROR: {e}")
            logger.error("[Groq] Full traceback:")
            import traceback
            logger.error(traceback.format_exc())
            return {"positions": [], "metadata": {"confidence": 0, "method": "groq_error"}}

    def _parse_with_regex(self, text: str) -> Dict:
        """Fallback regex parsing"""
        import re
        
        logger.info("[Regex] Starting regex parsing...")
        lines = text.split("\n")
        logger.info(f"[Regex] Total lines: {len(lines)}")
        
        positions = []
        
        for line in lines:
            # Format: "1 | Название | м | 140"
            match1 = re.match(r"^(\d+)\s*\|\s*(.*?)\s*\|\s*([а-яА-ЯмшШтТкКлЛ\s]+?)\s*\|\s*(\d+)", line)
            
            # Format: "1 Название м 140"
            match2 = re.match(r"^(\d+)\s+(.*?)\s+([а-яА-ЯмшШтТкКлЛ\s]+?)\s+(\d+)$", line)
            
            match = match1 or match2
            
            if match:
                pos_num = int(match.group(1))
                name = match.group(2).strip()
                unit = match.group(3).strip()
                qty = int(match.group(4))
                
                if name and unit:
                    positions.append({
                        "pos": pos_num,
                        "name": name,
                        "unit": unit,
                        "qty": qty
                    })
        
        logger.info(f"[Regex] Total positions found: {len(positions)}")
        
        return {
            "positions": positions,
            "metadata": {
                "confidence": 60,
                "method": "regex"
            },
            "preview": self._format_preview(positions),
            "raw_text": text
        }

    def _format_preview(self, positions: List[Dict]) -> str:
        """Format positions for preview display"""
        if not positions:
            return ""
        
        lines = []
        for item in positions:
            line = f"{item['pos']} | {item['name']} | {item['unit']} | {item['qty']}"
            lines.append(line)
        
        return "\n".join(lines)
