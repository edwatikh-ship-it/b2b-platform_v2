#!/usr/bin/env python3
"""
B2B Platform API - PRODUCTION VERSION
✅ Regex парсер + GROQ fallback с полным детектированием ошибок
"""

import os
import logging
import json
import re
from datetime import datetime
from pathlib import Path
from typing import List, Dict

# ✅ Загружаем .env переменные
from dotenv import load_dotenv
load_dotenv()

# ✅ FastAPI
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware

# ✅ Для работы с документами
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

try:
    from groq import Groq
    GROQ_AVAILABLE = True
except ImportError:
    GROQ_AVAILABLE = False

# ✅ ЛОГИРОВАНИЕ
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

logger.info(f"DOCX: {DOCX_AVAILABLE}, PDF: {PDF_AVAILABLE}, XLSX: {XLSX_AVAILABLE}, GROQ: {GROQ_AVAILABLE}")

# ✅ IN-MEMORY ХРАНИЛИЩЕ (ВМЕСТО БД)
requests_storage: Dict[int, dict] = {}
suppliers_storage: List[dict] = []
next_request_id = 1

def init_suppliers():
    """Инициализирует список поставщиков в памяти"""
    global suppliers_storage
    suppliers_storage = [
        {"id": 1, "name": "ООО Трубная компания", "inn": "7701234567", "url": "https://trubcom.ru", "rating": 4.8},
        {"id": 2, "name": "АО МеталлПродукт", "inn": "7702345678", "url": "https://metalprod.ru", "rating": 4.5},
        {"id": 3, "name": "ПАО Промышленные материалы", "inn": "7703456789", "url": "https://prommat.ru", "rating": 4.9},
        {"id": 4, "name": "ЛЛК Трубы и фитинги", "inn": "7704567890", "url": "https://llk-tubes.ru", "rating": 4.3},
        {"id": 5, "name": "ООО Сталь и конструкции", "inn": "7705678901", "url": "https://steel-constr.ru", "rating": 4.6},
        {"id": 6, "name": "АО Крепеж и фурнитура", "inn": "7706789012", "url": "https://hardware-pro.ru", "rating": 4.4},
        {"id": 7, "name": "ЗАО Промышленные решения", "inn": "7707890123", "url": "https://promsol.ru", "rating": 4.7},
    ]
    logger.info(f"✅ Initialized {len(suppliers_storage)} suppliers in memory")

# ✅ FastAPI app
app = FastAPI(
    title="B2B Platform API",
    version="0.2.0",
    description="API для B2B платформы (PRODUCTION)"
)

# ✅ CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ✅ ФУНКЦИИ ПАРСИНГА

def extract_text_from_file(file_path: str) -> str:
    """Извлекает текст из разных форматов файлов"""
    ext = file_path.lower().split('.')[-1]
    text = ""
    
    print(f"\n{'='*70}")
    print(f"🔍 EXTRACTING: {ext.upper()}")
    print(f"📁 FILE PATH: {file_path}")
    print(f"{'='*70}\n")
    logger.info(f"🔍 EXTRACTING: {ext.upper()}")
    
    try:
        if ext == "docx" and DOCX_AVAILABLE:
            doc = Document(file_path)
            
            para_count = len(doc.paragraphs)
            print(f"📄 PARAGRAPHS: {para_count} найдено")
            logger.info(f"📄 PARAGRAPHS: {para_count} найдено")
            
            paragraphs_text = []
            for i, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                if para_text:
                    paragraphs_text.append(para.text)
                    print(f"  ► [{i}] {para.text[:100]}")
            
            text = "\n".join(paragraphs_text)
            
            table_count = len(doc.tables)
            print(f"📊 TABLES: {table_count} найдено")
            logger.info(f"📊 TABLES: {table_count} найдено")
            
            tables_text = []
            for t_idx, table in enumerate(doc.tables):
                rows_count = len(table.rows)
                cols_count = len(table.rows[0].cells) if table.rows else 0
                print(f"  📋 TABLE #{t_idx}: {rows_count} rows x {cols_count} cols")
                
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            tables_text.append(cell_text)
            
            text = text + "\n" + "\n".join(tables_text)
            
        elif ext == "pdf" and PDF_AVAILABLE:
            with open(file_path, 'rb') as f:
                pdf = PyPDF2.PdfReader(f)
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
            
        elif ext == "xlsx" and XLSX_AVAILABLE:
            wb = openpyxl.load_workbook(file_path)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                for row in ws.iter_rows(values_only=True):
                    row_text = " ".join([str(cell).strip() if cell else "" for cell in row if cell])
                    if row_text.strip():
                        text += row_text + "\n"
            
        elif ext == "txt":
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
    
    except Exception as e:
        logger.error(f"❌ ERROR EXTRACTING: {e}", exc_info=True)
    
    final_text = text.strip()
    print(f"\n{'='*70}")
    print(f"🎯 FINAL TEXT LENGTH: {len(final_text)} chars")
    print(f"{'='*70}\n")
    
    return final_text

def parse_text_regex(text: str) -> dict:
    """РАБОТАЮЩИЙ парсер - строго по структуре таблицы"""
    
    lines = text.split('\n')
    items = []
    
    print(f"\n{'='*70}")
    print(f"🔧 PARSING (REGEX): {len(lines)} lines")
    print(f"{'='*70}\n")
    logger.info(f"🔧 PARSING (REGEX): {len(lines)} lines")
    
    table_start = -1
    for i, line in enumerate(lines):
        if line.strip() == '№':
            table_start = i
            print(f"✅ TABLE STARTS at line {i}")
            logger.info(f"✅ TABLE STARTS at line {i}")
            break
    
    if table_start == -1:
        print(f"❌ NO TABLE FOUND")
        logger.warning("❌ NO TABLE FOUND")
        return {"positions": [], "confidence": 0, "source": "regex"}
    
    i = table_start + 1
    
    while i < len(lines):
        line = lines[i].strip()
        
        if line.isdigit() and len(line) == 1:
            pos_num = int(line)
            print(f"\n📍 Position #{pos_num}")
            logger.info(f"📍 Position #{pos_num}")
            
            if i + 3 >= len(lines):
                print(f"⚠️  Not enough lines for position {pos_num}")
                break
            
            name_raw = lines[i + 1].strip()
            unit_raw = lines[i + 2].strip()
            qty_raw = lines[i + 3].strip()
            
            print(f"  Name:  {name_raw}")
            print(f"  Unit:  {unit_raw}")
            print(f"  Qty:   {qty_raw}")
            
            name = name_raw
            name = re.sub(r'\s+\([^)]*\)', '', name)
            name = re.sub(r'\s+', ' ', name).strip()
            
            unit = unit_raw if unit_raw in ['м', 'м²', 'м³', 'шт', 'кг', 'т', 'л', 'см', 'мм'] else 'шт'
            
            qty_match = re.search(r'(\d+)', qty_raw)
            qty = int(qty_match.group(1)) if qty_match else 1
            
            print(f"  → Cleaned: {name} ({qty} {unit})")
            
            items.append({
                'pos': len(items) + 1,
                'name': name,
                'qty': qty,
                'unit': unit
            })
            
            logger.info(f"  ✅ Added: {name} ({qty} {unit})")
            i += 4
        else:
            i += 1
    
    print(f"\n{'='*70}")
    print(f"📊 RESULT: {len(items)} items (source: regex)")
    for item in items:
        print(f"  [{item['pos']}] {item['name']} ({item['qty']} {item['unit']})")
    print(f"{'='*70}\n")
    
    logger.info(f"📊 REGEX RESULT: {len(items)} items")
    
    confidence = 85 if len(items) > 0 else 0
    return {
        "positions": items,
        "confidence": confidence,
        "source": "regex"
    }

def parse_text_with_groq(text: str) -> dict:
    """
    Парсит текст с GROQ с полным детектированием падений
    Автоматически переходит на regex если GROQ не работает
    """
    if not text or not text.strip():
        logger.warning("⚠️  Empty text provided")
        return {"positions": [], "confidence": 0, "source": "empty"}
    
    if not GROQ_AVAILABLE or not os.getenv("GROQ_API_KEY"):
        logger.warning("⚠️  GROQ not available, using regex")
        return parse_text_regex(text)
    
    try:
        logger.info("🔄 Trying GROQ...")
        client = Groq(api_key=os.getenv("GROQ_API_KEY"))
        
        system_prompt = """Ты - парсер закупок. Извлеки позиции товаров.
Верни ТОЛЬКО JSON:
[{"pos": 1, "name": "Товар", "unit": "м", "qty": 100}]"""
        
        chat_completion = client.chat.completions.create(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"{text[:8000]}"}
            ],
            model="llama-3.1-70b-versatile",
            temperature=0.1,
            max_tokens=2048,
            timeout=15
        )
        
        response_text = chat_completion.choices[0].message.content.strip()
        logger.info(f"✅ GROQ response received: {len(response_text)} chars")
        
        json_text = response_text
        if "```json" in response_text:
            json_text = response_text.split("```json")[1].split("```")[0].strip()
        elif "```" in response_text:
            json_text = response_text.split("```")[1].split("```")[0].strip()
        
        items = json.loads(json_text)
        
        if isinstance(items, list) and len(items) > 0:
            logger.info(f"✅ GROQ SUCCESS: {len(items)} items, confidence=90%")
            return {"positions": items, "confidence": 90, "source": "groq"}
        else:
            logger.warning(f"⚠️  GROQ returned empty list")
            return parse_text_regex(text)
    
    except Exception as e:
        error_str = str(e).lower()
        
        if "decommissioned" in error_str or ("model" in error_str and "no longer" in error_str):
            logger.error(f"❌ GROQ MODEL DECOMMISSIONED: {e}")
            logger.warning("⚠️  Model no longer supported, using regex")
        elif "429" in error_str or "rate" in error_str or "limit" in error_str:
            logger.error(f"❌ GROQ RATE LIMIT: {e}")
            logger.warning("⚠️  Rate limit exceeded, using regex")
        elif "timeout" in error_str or "time out" in error_str:
            logger.error(f"❌ GROQ TIMEOUT: {e}")
            logger.warning("⚠️  Request timeout, using regex")
        elif "500" in error_str or ("server" in error_str and "error" in error_str):
            logger.error(f"❌ GROQ SERVER ERROR: {e}")
            logger.warning("⚠️  Server error, using regex")
        elif "connection" in error_str or "network" in error_str or "refused" in error_str:
            logger.error(f"❌ GROQ CONNECTION ERROR: {e}")
            logger.warning("⚠️  Connection error, using regex")
        elif "401" in error_str or "unauthorized" in error_str or ("invalid" in error_str and "key" in error_str):
            logger.error(f"❌ GROQ AUTH ERROR: {e}")
            logger.warning("⚠️  Invalid API key, using regex")
        else:
            logger.error(f"❌ GROQ UNKNOWN ERROR: {e}", exc_info=True)
            logger.warning("⚠️  Unknown error, using regex")
        
        return parse_text_regex(text)

# ✅ API ENDPOINTS

@app.post("/api/v1/user/upload-and-create")
async def upload_and_create(file: UploadFile = File(...)):
    """Загрузить файл"""
    global next_request_id
    
    try:
        logger.info(f"UPLOAD: {file.filename}")
        
        upload_dir = Path(os.getcwd()) / "uploads"
        upload_dir.mkdir(exist_ok=True)
        
        file_path = upload_dir / f"{datetime.now().timestamp()}_{file.filename}"
        
        content = await file.read()
        with open(file_path, "wb") as f:
            f.write(content)
        
        logger.info(f"FILE SAVED: {file_path}")
        
        request_id = next_request_id
        next_request_id += 1
        
        requests_storage[request_id] = {
            "id": request_id,
            "filename": file.filename,
            "status": "draft",
            "created_at": datetime.now().isoformat(),
            "file_path": str(file_path),
            "items": [],
            "parsing_confidence": 0,
            "preview": "",
            "parsing_source": "unknown"
        }
        
        logger.info(f"REQUEST CREATED: #{request_id}")
        
        return {"success": True, "request_id": request_id, "filename": file.filename}
    
    except Exception as e:
        logger.error(f"ERROR UPLOAD: {e}")
        raise HTTPException(status_code=400, detail=str(e))

@app.get("/api/v1/user/requests")
async def get_user_requests():
    """Получить все заявки"""
    logger.info(f"GET REQUESTS: {len(requests_storage)} total")
    
    return [
        {
            "id": r["id"],
            "filename": r["filename"],
            "status": r["status"],
            "created_at": r["created_at"],
            "items_count": len(r["items"]),
            "confidence": r["parsing_confidence"],
            "preview": r["preview"][:100] if r["preview"] else ""
        }
        for r in requests_storage.values()
    ]

@app.get("/api/v1/user/requests/{request_id}")
async def get_request_detail(request_id: int):
    """Получить деталь заявки"""
    
    if request_id not in requests_storage:
        raise HTTPException(status_code=404, detail="Request not found")
    
    r = requests_storage[request_id]
    return {
        "id": r["id"],
        "filename": r["filename"],
        "status": r["status"],
        "created_at": r["created_at"],
        "items": r["items"],
        "confidence": r["parsing_confidence"],
        "preview": r["preview"],
        "parsing_source": r["parsing_source"]
    }

@app.post("/api/v1/user/requests/{request_id}/submit")
async def submit_request(request_id: int):
    """Отправить на распознавание"""
    
    logger.info("=" * 60)
    logger.info(f"PARSING START: REQUEST #{request_id}")
    logger.info("=" * 60)
    
    if request_id not in requests_storage:
        raise HTTPException(status_code=404, detail="Request not found")
    
    try:
        r = requests_storage[request_id]
        
        if r["status"] != "draft":
            raise HTTPException(status_code=400, detail="Can only submit draft requests")
        
        if not os.path.exists(r["file_path"]):
            raise HTTPException(status_code=400, detail="File not found")
        
        text = extract_text_from_file(r["file_path"])
        logger.info(f"TEXT EXTRACTED: {len(text)} chars")
        
        parse_result = parse_text_with_groq(text)
        
        items = parse_result.get("positions", [])
        confidence = parse_result.get("confidence", 0)
        source = parse_result.get("source", "unknown")
        
        r["status"] = "submitted"
        r["items"] = items
        r["parsing_confidence"] = confidence
        r["parsing_source"] = source
        r["preview"] = text[:500]
        
        logger.info(f"REQUEST UPDATED: {len(items)} items, confidence={confidence}%, source={source}")
        logger.info("=" * 60)
        
        return {
            "success": True,
            "items_count": len(items),
            "confidence": confidence,
            "preview": r["preview"],
            "message": f"Found {len(items)} positions"
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"PARSING ERROR: {e}")
        import traceback
        traceback.print_exc()
        logger.info("=" * 60)
        raise HTTPException(status_code=400, detail=str(e))

@app.get("/api/v1/suppliers/search")
async def search_suppliers(keyword: str = ""):
    """Поиск поставщиков"""
    logger.info(f"SEARCH SUPPLIERS: '{keyword}'")
    
    if not keyword:
        results = suppliers_storage
    else:
        keyword_lower = keyword.lower()
        results = [s for s in suppliers_storage if keyword_lower in s["name"].lower()]
    
    logger.info(f"FOUND: {len(results)} suppliers")
    
    return {
        "results": [
            {
                "id": s["id"],
                "name": s["name"],
                "inn": s["inn"],
                "url": s["url"],
                "rating": s["rating"]
            }
            for s in results
        ]
    }

@app.get("/api/v1/moderator/tasks")
async def get_moderator_tasks():
    """Получить список задач на модерацию"""
    tasks = [
        {
            "id": r["id"],
            "request_id": r["id"],
            "filename": r["filename"],
            "status": "pending",
            "parsing_method": r["parsing_source"],
            "items_count": len(r["items"]),
            "confidence": r["parsing_confidence"],
            "created_at": r["created_at"]
        }
        for r in requests_storage.values()
        if r["status"] == "submitted"
    ]
    
    logger.info(f"MODERATOR TASKS: {len(tasks)}")
    return tasks

@app.get("/api/v1/moderator/tasks/{task_id}")
async def get_task_detail(task_id: int):
    """Получить деталь задачи"""
    
    if task_id not in requests_storage:
        raise HTTPException(status_code=404, detail="Task not found")
    
    r = requests_storage[task_id]
    
    return {
        "id": task_id,
        "request_id": task_id,
        "filename": r["filename"],
        "status": "pending",
        "parsing_method": r["parsing_source"],
        "confidence": r["parsing_confidence"],
        "items": r["items"],
        "created_at": r["created_at"],
        "preview": r["preview"]
    }

@app.post("/api/v1/moderator/tasks/{task_id}/approve")
async def approve_task(task_id: int):
    """Одобрить задачу"""
    
    if task_id not in requests_storage:
        raise HTTPException(status_code=404, detail="Task not found")
    
    requests_storage[task_id]["status"] = "approved"
    logger.info(f"TASK APPROVED: #{task_id}")
    
    return {"success": True, "message": f"Task #{task_id} approved"}

@app.post("/api/v1/moderator/tasks/{task_id}/reject")
async def reject_task(task_id: int):
    """Отклонить задачу"""
    
    if task_id not in requests_storage:
        raise HTTPException(status_code=404, detail="Task not found")
    
    requests_storage[task_id]["status"] = "rejected"
    logger.info(f"TASK REJECTED: #{task_id}")
    
    return {"success": True, "message": f"Task #{task_id} rejected"}

@app.get("/health")
async def health_check():
    return {"status": "ok", "version": "0.2.0"}

@app.get("/")
async def root():
    return {"name": "B2B Platform API", "version": "0.2.0", "status": "working"}

@app.on_event("startup")
async def startup_event():
    init_suppliers()

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)
